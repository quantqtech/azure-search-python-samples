"""
Convert the ROI Model markdown to a professional Word document.
Uses python-docx to create formatted .docx with title page, tables, and styling.
"""
import sys
import os

# Fix Windows encoding for stdout
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import re

# ── Configuration ──
OUTPUT_PATH = r"c:\repos\my-projects\azure-search-python-samples\Gent-Project-Management\ROI Model - Gent Davenport.docx"
NAVY = RGBColor(0x1B, 0x2A, 0x4A)  # dark navy for headers
LIGHT_GRAY = "D9E2F3"  # alternating row shading
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
BLACK = RGBColor(0x00, 0x00, 0x00)
DARK_BLUE_HEX = "1B2A4A"
ACCENT_BLUE = RGBColor(0x2E, 0x75, 0xB6)

doc = Document()

# ── Set default font to Calibri 11pt ──
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
font.color.rgb = BLACK

# Set heading styles to Calibri
for i in range(1, 4):
    heading_style = doc.styles[f'Heading {i}']
    heading_style.font.name = 'Calibri'
    heading_style.font.color.rgb = NAVY

doc.styles['Heading 1'].font.size = Pt(22)
doc.styles['Heading 1'].font.bold = True
doc.styles['Heading 2'].font.size = Pt(16)
doc.styles['Heading 2'].font.bold = True
doc.styles['Heading 3'].font.size = Pt(13)
doc.styles['Heading 3'].font.bold = True

# ── Set 1-inch margins on all sections ──
for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)


# ── Helper Functions ──

def add_page_numbers(doc):
    """Add page numbers to footer."""
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # PAGE field
        run = p.add_run()
        run.font.name = 'Calibri'
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
        fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
        run._r.append(fldChar1)
        run2 = p.add_run()
        run2.font.name = 'Calibri'
        run2.font.size = Pt(9)
        instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
        run2._r.append(instrText)
        run3 = p.add_run()
        run3.font.name = 'Calibri'
        run3.font.size = Pt(9)
        fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
        run3._r.append(fldChar2)


def set_cell_shading(cell, color_hex):
    """Apply background shading to a table cell."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_borders(table):
    """Add borders to all cells in a table."""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="4" w:space="0" w:color="BFBFBF"/>'
        f'  <w:left w:val="single" w:sz="4" w:space="0" w:color="BFBFBF"/>'
        f'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="BFBFBF"/>'
        f'  <w:right w:val="single" w:sz="4" w:space="0" w:color="BFBFBF"/>'
        f'  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="BFBFBF"/>'
        f'  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="BFBFBF"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)


def add_formatted_text(paragraph, text):
    """Parse markdown bold (**text**) and add runs with appropriate formatting."""
    # Split on **...**  patterns
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
            run.font.name = 'Calibri'
            run.font.size = Pt(11)
        else:
            run = paragraph.add_run(part)
            run.font.name = 'Calibri'
            run.font.size = Pt(11)


def add_formatted_cell_text(cell, text, bold=False, font_size=Pt(10)):
    """Add formatted text to a table cell, handling markdown bold."""
    cell.text = ""  # clear default
    p = cell.paragraphs[0]
    p.space_before = Pt(2)
    p.space_after = Pt(2)

    if bold:
        run = p.add_run(text.replace('**', ''))
        run.bold = True
        run.font.name = 'Calibri'
        run.font.size = font_size
        return

    # Handle **bold** within text
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            run.bold = True
            run.font.name = 'Calibri'
            run.font.size = font_size
        else:
            run = p.add_run(part)
            run.font.name = 'Calibri'
            run.font.size = font_size


def create_table(doc, headers, rows):
    """Create a professionally formatted table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    set_cell_borders(table)

    # Header row — navy background, white text
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, DARK_BLUE_HEX)
        add_formatted_cell_text(cell, header.replace('**', ''), bold=True, font_size=Pt(10))
        for run in cell.paragraphs[0].runs:
            run.font.color.rgb = WHITE

    # Data rows — alternating shading
    for row_idx, row_data in enumerate(rows):
        for col_idx, cell_text in enumerate(row_data):
            cell = table.rows[row_idx + 1].cells[col_idx]
            add_formatted_cell_text(cell, cell_text, font_size=Pt(10))
            # Alternating row shading
            if row_idx % 2 == 1:
                set_cell_shading(cell, LIGHT_GRAY)

    # Add spacing after table
    p = doc.add_paragraph()
    p.space_before = Pt(4)
    p.space_after = Pt(4)

    return table


def add_paragraph_with_formatting(doc, text, style=None):
    """Add a paragraph with markdown bold support."""
    p = doc.add_paragraph(style=style)
    add_formatted_text(p, text)
    return p


def add_bullet(doc, text):
    """Add a bullet point with markdown bold support."""
    p = doc.add_paragraph(style='List Bullet')
    # Clear default text and use formatted text
    add_formatted_text(p, text)
    for run in p.runs:
        run.font.name = 'Calibri'
        run.font.size = Pt(11)
    return p


# ══════════════════════════════════════════
# TITLE PAGE
# ══════════════════════════════════════════

# Add some vertical spacing
for _ in range(6):
    doc.add_paragraph()

# Title
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("ROI Model \u2014 Gent Machine")
run.font.name = 'Calibri'
run.font.size = Pt(32)
run.font.color.rgb = NAVY
run.bold = True

# Subtitle
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("The Growth Case for Phase 2")
run.font.name = 'Calibri'
run.font.size = Pt(18)
run.font.color.rgb = ACCENT_BLUE

# Horizontal line (using a thin table as a visual separator)
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("_" * 60)
run.font.color.rgb = RGBColor(0xBF, 0xBF, 0xBF)
run.font.size = Pt(8)

doc.add_paragraph()

# Prepared by
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Prepared by: Aidan Systems")
run.font.name = 'Calibri'
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x40, 0x40, 0x40)

# Date
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Date: March 2026")
run.font.name = 'Calibri'
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x40, 0x40, 0x40)

# For
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("For: Gent Machine leadership (via MAGNET)")
run.font.name = 'Calibri'
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x40, 0x40, 0x40)

# Page break after title page
doc.add_page_break()

# ══════════════════════════════════════════
# SECTION 1: THE SITUATION
# ══════════════════════════════════════════

doc.add_heading('The Situation', level=1)

p = doc.add_paragraph()
add_formatted_text(p, "Gent Machine has a large fleet of Davenport screw machines \u2014 their website lists 38. Not all are running at capacity. The constraint isn't equipment \u2014 ")
run = p.add_run("it's skilled people.")
run.bold = True
run.font.name = 'Calibri'
run.font.size = Pt(11)
add_formatted_text(p, " Gent's longtime expert (Dave) has retired. The knowledge to set up, troubleshoot, and maintain Davenports at production quality lives in whatever the current team learned from Dave, plus the manuals on the shelf.")

p = doc.add_paragraph()
run = p.add_run("The core question: could Gent get more Davenports running if the skills bottleneck were removed?")
run.bold = True
run.font.name = 'Calibri'
run.font.size = Pt(11)
add_formatted_text(p, " Even one additional machine represents significant revenue.")

p = doc.add_paragraph()
add_formatted_text(p, "Phase 1 captured Dave's expertise and Gent's maintenance documentation into an AI-powered maintenance assistant. It's live on the shop floor. Machinists have used it for real troubleshooting.")

p = doc.add_paragraph()
run = p.add_run("Phase 2 isn't about protecting what you have. It's about growing what you can do.")
run.bold = True
run.font.name = 'Calibri'
run.font.size = Pt(11)

doc.add_page_break()

# ══════════════════════════════════════════
# SECTION 2: THE OPPORTUNITY
# ══════════════════════════════════════════

doc.add_heading("The Opportunity: What's One More Davenport Worth?", level=1)

doc.add_heading('Revenue Per Machine', level=2)

create_table(doc,
    ["Input", "Estimate", "Notes"],
    [
        ["Shop billing rate (multi-spindle screw machine)", "$85-$100/hr", "Industry range for Davenports; east coast shops bill ~$100/hr"],
        ["Productive hours per day", "6-7 hrs", "Accounting for setup, changeover, downtime"],
        ["Operating days per month", "20", "Standard single shift"],
        ["**Revenue per machine per month**", "**$10,200 - $14,000**", ""],
        ["**Revenue per machine per year**", "**$122,400 - $168,000**", ""],
    ]
)

p = doc.add_paragraph()
p.style = doc.styles['Normal']
run = p.add_run("These are billing rates, not margin. Gent's actual margin depends on their cost structure \u2014 labor, materials, overhead. But even at 30-40% margin, one more Davenport running is $36,000 - $67,000/year in additional profit.")
run.italic = True
run.font.name = 'Calibri'
run.font.size = Pt(11)

doc.add_heading("The Constraint: Why Can't Gent Run More Machines?", level=2)

create_table(doc,
    ["Constraint", "Without AI Cockpit", "With AI Cockpit"],
    [
        ["**Setup knowledge**", "Takes months/years to learn Davenport setup. Dave used to teach this. Dave is retired.", "System provides step-by-step setup guidance with citations from manuals and Dave's captured expertise"],
        ["**Troubleshooting**", "When something goes wrong, junior machinists guess, wait, or call for help. Extended downtime.", "30-second diagnosis with ranked causes and fix procedures"],
        ["**Hiring pool**", "Must hire experienced Davenport machinists \u2014 rare, expensive, competing with other shops", "Can hire trainable people and pair them with the cockpit. The system bridges the experience gap."],
        ["**Training time**", "6-12+ months to independently troubleshoot (and that was WITH Dave mentoring)", "Target: productive in weeks, independently troubleshooting common issues in 2-3 months"],
        ["**Confidence**", "Management won't put a junior person on a machine alone \u2014 risk of scrap, damage, downtime", "System provides safety net \u2014 guidance, checklists, escalation paths"],
    ]
)

doc.add_page_break()

# ══════════════════════════════════════════
# SECTION 3: THE VALUE MODEL
# ══════════════════════════════════════════

doc.add_heading('The Value Model: Scale, Capability, Redundancy', level=1)

# ── 1. SCALE ──
doc.add_heading('1. SCALE \u2014 Get More Machines Running', level=2)

p = doc.add_paragraph()
run = p.add_run("The headline number.")
run.bold = True
run.font.name = 'Calibri'
run.font.size = Pt(11)
add_formatted_text(p, " Every idle Davenport is lost revenue.")

create_table(doc,
    ["Scenario", "New Machines", "Annual Revenue Added", "Annual Margin (est. 35%)"],
    [
        ["**Conservative** \u2014 1 more Davenport", "+1", "$120,000 - $168,000", "**$42,000 - $59,000**"],
        ["**Moderate** \u2014 2 more Davenports", "+2", "$240,000 - $336,000", "**$84,000 - $118,000**"],
    ]
)

p = doc.add_paragraph()
run = p.add_run("Gent knows better than anyone how many additional machines are viable and what the demand looks like. The model only needs ONE more running to justify the investment many times over.")
run.italic = True
run.font.name = 'Calibri'
run.font.size = Pt(11)

p = doc.add_paragraph()
add_formatted_text(p, "Phase 2 Managed AIR cost to Gent: **~$15,600/year** (Year 1) dropping to **~$12,500/year** ongoing.")

p = doc.add_paragraph()
run = p.add_run("Getting ONE more Davenport running pays for 3-4 years of Managed AIR in the first year alone.")
run.bold = True
run.font.name = 'Calibri'
run.font.size = Pt(11)

p = doc.add_paragraph("What it takes to get there:")
add_bullet(doc, "Hire 1 trainable machinist (doesn't need to be an experienced Davenport operator)")
add_bullet(doc, "Pair them with the AI cockpit from day one")
add_bullet(doc, "Phase 2 Managed AIR expands the knowledge base to cover setup procedures, not just troubleshooting")
add_bullet(doc, "Flight Instructor (Aidan) coaches the onboarding and monitors the new operator's query patterns to identify knowledge gaps in real-time")

# ── 2. CAPABILITY ──
doc.add_heading('2. CAPABILITY \u2014 Raise the Floor Across All Operators', level=2)

p = doc.add_paragraph()
add_formatted_text(p, "Every machinist on the floor \u2014 experienced or junior \u2014 gets access to expert-level guidance. This doesn't just help new hires. It makes existing operators better.")

create_table(doc,
    ["Benefit", "Impact", "How to Measure"],
    [
        ["**Junior operators troubleshoot independently**", "Fewer escalations, less downtime waiting for help", "Track: queries resolved without escalation"],
        ["**Consistent quality across machines**", "Fewer scrap incidents from incorrect setup or adjustment", "Track: scrap rate before/after"],
        ["**Faster changeover**", "Setup procedures available instantly vs. digging through manuals", "Track: changeover time"],
        ["**Shift independence**", "Every shift has equal troubleshooting capability, not dependent on who's working", "Track: downtime by shift"],
    ]
)

p = doc.add_paragraph()
run = p.add_run("The capability gap matters because Gent serves demanding industries")
run.bold = True
run.font.name = 'Calibri'
run.font.size = Pt(11)
add_formatted_text(p, " \u2014 automotive (IATF 16949 certified), defense, hydraulics. Quality isn't optional. The AI cockpit helps junior operators meet the same quality standard that experienced operators hit.")

# ── 3. REDUNDANCY ──
doc.add_heading('3. REDUNDANCY \u2014 Remove Single Points of Failure', level=2)

p = doc.add_paragraph()
add_formatted_text(p, "With 8 machines and a small team, one person being out sick can impact production. Dave's retirement already proved what happens when the single expert leaves.")

create_table(doc,
    ["Risk", "Without System", "With System"],
    [
        ["Experienced operator calls in sick", "That machine may not run, or runs at reduced capability", "Junior operator + cockpit can cover"],
        ["Second experienced operator leaves", "Back to the Dave problem", "Knowledge is in the system, not just in heads"],
        ["Vacation coverage", "Limited \u2014 only certain people can run certain setups", "Any trained operator + cockpit can handle standard setups"],
        ["Customer surge / rush order", "Can't add capacity because can't staff more machines", "Flex up \u2014 put a cockpit-equipped operator on an idle machine"],
    ]
)

doc.add_page_break()

# ══════════════════════════════════════════
# SECTION 4: THE GROWTH PATH
# ══════════════════════════════════════════

doc.add_heading('The Growth Path: How Phase 2 Gets Gent to Machine #9', level=1)

create_table(doc,
    ["Month", "Stage", "What's Happening", "Milestone"],
    [
        ["1", "TAXI / Train & Tune", "Complete structured training for all current machinists. Close top knowledge gaps.", "100% of team has used the cockpit"],
        ["2-3", "WHEELS UP / Adopt", "\u201cAsk the Cockpit First\u201d is the team norm. Feedback flowing. KB expanding.", "Regular daily use across the team"],
        ["3-4", "WHEELS UP", "Begin adding setup procedures to the KB \u2014 not just troubleshooting. This is what enables a junior operator to set up a machine.", "Setup content for top 5 job types indexed"],
        ["4-5", "CLIMBING / Expand", "Hire junior machinist. Pair with cockpit from day one. Flight Instructor monitors their queries and closes gaps in real-time.", "New hire productive on basic jobs"],
        ["5-6", "CLIMBING", "New hire handling standard setups and troubleshooting with cockpit support. Experienced operators freed up for complex work.", "Machine #9 running on basic jobs"],
        ["7+", "CRUISING / Compound", "System is comprehensive. New hire is independent for common scenarios. Knowledge base covers setup + troubleshooting + maintenance.", "Machine #9 running at full production"],
    ]
)

p = doc.add_paragraph()
run = p.add_run("Timeline to Machine #9: approximately 6 months from Phase 2 kickoff.")
run.bold = True
run.font.name = 'Calibri'
run.font.size = Pt(11)

doc.add_page_break()

# ══════════════════════════════════════════
# SECTION 5: INVESTMENT VS. RETURN
# ══════════════════════════════════════════

doc.add_heading('Investment vs. Return', level=1)

# Year 1
doc.add_heading('Year 1', level=2)

create_table(doc,
    ["", "Cost", "Revenue/Value"],
    [
        ["Phase 2 Managed AIR (6 mo Medium + 6 mo Small)", "~$15,600", ""],
        ["Junior machinist hire (fully loaded, 8 months)", "~$40,000", ""],
        ["**Total investment**", "**~$55,600**", ""],
        ["", "", ""],
        ["Revenue from Machine #9 (conservative, 6 months running)", "", "~$60,000 - $84,000"],
        ["Capability improvement across existing 8 machines (fewer escalations, less downtime, less scrap \u2014 est. 5% efficiency gain)", "", "~$10,000 - $15,000"],
        ["**Total Year 1 value**", "", "**~$70,000 - $99,000**"],
        ["**Net Year 1 value**", "", "**$14,400 - $43,400**"],
    ]
)

# Year 2
doc.add_heading('Year 2 (Machine #9 running full year, Small tier)', level=2)

create_table(doc,
    ["", "Cost", "Revenue/Value"],
    [
        ["Managed AIR (Small tier)", "~$12,500", ""],
        ["Junior machinist (full year)", "~$60,000", ""],
        ["**Total cost**", "**~$72,500**", ""],
        ["Revenue from Machine #9 (full year)", "", "~$120,000 - $168,000"],
        ["Capability improvement across 9 machines", "", "~$12,000 - $18,000"],
        ["**Total Year 2 value**", "", "**~$132,000 - $186,000**"],
        ["**Net Year 2 value**", "", "**$59,500 - $113,500**"],
    ]
)

# Year 2+
doc.add_heading('Year 2+ with Machine #10', level=2)

p = doc.add_paragraph()
add_formatted_text(p, "Add another $10,000 (work center expansion for new job types) + another junior hire. The platform is already running. Each incremental machine is cheaper to enable.")

doc.add_page_break()

# ══════════════════════════════════════════
# SECTION 6: THE CASE FOR MAGNET / MEP
# ══════════════════════════════════════════

doc.add_heading('The Case for MAGNET / MEP Impact', level=1)

create_table(doc,
    ["MEP Metric", "Phase 2 Impact"],
    [
        ["**New sales**", "$120,000-$168,000/year per additional Davenport brought online"],
        ["**Jobs created**", "1-2 junior machinists hired to run previously idle machines"],
        ["**New investment**", "Gent self-funds AI operations ($15,600 Year 1) + new hire"],
        ["**Cost savings**", "5%+ efficiency improvement across existing 8 machines"],
        ["**Jobs retained**", "Knowledge preservation \u2014 Dave's retirement didn't kill capability"],
    ]
)

doc.add_page_break()

# ══════════════════════════════════════════
# SECTION 7: ONE-PAGE SUMMARY
# ══════════════════════════════════════════

doc.add_heading('What Gent Sees: The One-Page Summary', level=1)

p = doc.add_paragraph()
run = p.add_run("You have 38 Davenports. 8 are running. The bottleneck is knowledge, not equipment.")
run.bold = True
run.font.name = 'Calibri'
run.font.size = Pt(11)

p = doc.add_paragraph()
add_formatted_text(p, "Phase 1 captured your expert's 30 years of knowledge into an AI system. Your machinists have already used it to troubleshoot real problems on the shop floor.")

p = doc.add_paragraph()
add_formatted_text(p, "Phase 2 turns that knowledge into growth:")

create_table(doc,
    ["", "What It Does", "What It's Worth"],
    [
        ["**Scale**", "Get another Davenport running by pairing a junior hire with AI-guided troubleshooting and setup", "$120,000-$168,000/year in new production capacity"],
        ["**Capability**", "Every machinist on the floor gets expert-level guidance \u2014 better quality, faster changeover, fewer mistakes", "5%+ efficiency across existing 8 machines"],
        ["**Redundancy**", "No single person is a bottleneck anymore \u2014 any shift, any operator can handle standard issues", "Risk reduction \u2014 one person out doesn't stop a machine"],
    ]
)

p = doc.add_paragraph()
run = p.add_run("The cost: ")
run.bold = True
run.font.name = 'Calibri'
run.font.size = Pt(11)
add_formatted_text(p, "~$1,560/month for 6 months, then ~$1,040/month ongoing. Plus a junior hire you were going to need anyway.")

p = doc.add_paragraph()
run = p.add_run("The math: ")
run.bold = True
run.font.name = 'Calibri'
run.font.size = Pt(11)
add_formatted_text(p, "Getting one more Davenport running pays for Phase 2 four times over.")

p = doc.add_paragraph()
run = p.add_run("The question isn't whether you can afford Phase 2. It's how many more machines could be running if the skills bottleneck were removed \u2014 and the knowledge to do that is already in the system.")
run.bold = True
run.font.name = 'Calibri'
run.font.size = Pt(11)

# ══════════════════════════════════════════
# Add page numbers and save
# ══════════════════════════════════════════

add_page_numbers(doc)
doc.save(OUTPUT_PATH)
print(f"Document saved to: {OUTPUT_PATH}")
