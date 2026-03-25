"""
Generate Phase 2 SOW Word document from markdown content.
Uses python-docx to create a professionally formatted .docx file.
"""

import sys
import os

# Fix Windows UTF-8 stdout
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import re

# Colors
NAVY = RGBColor(0x1B, 0x2A, 0x4A)
LIGHT_GRAY = RGBColor(0xF2, 0xF2, 0xF2)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
BLACK = RGBColor(0x00, 0x00, 0x00)
DARK_GRAY = RGBColor(0x33, 0x33, 0x33)

OUTPUT_PATH = os.path.join(
    r"c:\repos\my-projects\azure-search-python-samples\Gent-Project-Management",
    "Phase 2 SOW - Managed AIR - Gent.docx"
)


def set_cell_shading(cell, color_hex):
    """Apply background shading to a table cell."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_font(run, name="Calibri", size=11, bold=False, color=None, italic=False):
    """Set font properties on a run."""
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color


def add_formatted_text(paragraph, text, default_size=11, default_color=BLACK):
    """Parse text with **bold** markers and add runs accordingly."""
    # Split on **bold** patterns
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            set_font(run, size=default_size, bold=True, color=default_color)
        elif part.startswith('*') and part.endswith('*'):
            run = paragraph.add_run(part[1:-1])
            set_font(run, size=default_size, italic=True, color=default_color)
        else:
            run = paragraph.add_run(part)
            set_font(run, size=default_size, color=default_color)


def set_table_borders(table):
    """Add borders to all cells in a table."""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '  <w:top w:val="single" w:sz="4" w:space="0" w:color="BFBFBF"/>'
        '  <w:left w:val="single" w:sz="4" w:space="0" w:color="BFBFBF"/>'
        '  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="BFBFBF"/>'
        '  <w:right w:val="single" w:sz="4" w:space="0" w:color="BFBFBF"/>'
        '  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="BFBFBF"/>'
        '  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="BFBFBF"/>'
        '</w:tblBorders>'
    )
    tblPr.append(borders)


def add_table(doc, headers, rows):
    """Add a formatted table to the document."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    set_table_borders(table)

    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        add_formatted_text(p, header, default_size=10, default_color=WHITE)
        for run in p.runs:
            run.bold = True
        set_cell_shading(cell, "1B2A4A")

    # Data rows
    for row_idx, row_data in enumerate(rows):
        for col_idx, cell_text in enumerate(row_data):
            cell = table.rows[row_idx + 1].cells[col_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            add_formatted_text(p, cell_text, default_size=10, default_color=DARK_GRAY)
            # Alternating row shading
            if row_idx % 2 == 1:
                set_cell_shading(cell, "F2F2F2")

    # Set cell padding
    for row in table.rows:
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcMar = parse_xml(
                f'<w:tcMar {nsdecls("w")}>'
                '  <w:top w:w="40" w:type="dxa"/>'
                '  <w:left w:w="80" w:type="dxa"/>'
                '  <w:bottom w:w="40" w:type="dxa"/>'
                '  <w:right w:w="80" w:type="dxa"/>'
                '</w:tcMar>'
            )
            tcPr.append(tcMar)

    doc.add_paragraph()  # spacing after table
    return table


def add_page_number(doc):
    """Add page numbers to the footer."""
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add PAGE field
        run = p.add_run()
        set_font(run, size=9, color=RGBColor(0x80, 0x80, 0x80))

        fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
        run._r.append(fldChar1)

        run2 = p.add_run()
        set_font(run2, size=9, color=RGBColor(0x80, 0x80, 0x80))
        instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
        run2._r.append(instrText)

        run3 = p.add_run()
        set_font(run3, size=9, color=RGBColor(0x80, 0x80, 0x80))
        fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
        run3._r.append(fldChar2)


def create_title_page(doc):
    """Create the title page."""
    # Add some blank paragraphs for vertical centering
    for _ in range(6):
        p = doc.add_paragraph()
        p.space_after = Pt(0)
        p.space_before = Pt(0)

    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Phase 2: Managed AIR")
    set_font(run, size=28, bold=True, color=NAVY)

    # Subtitle line 1
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.space_before = Pt(6)
    run = p.add_run("Statement of Work")
    set_font(run, size=22, color=NAVY)

    # Horizontal rule (thin line)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.space_before = Pt(24)
    p.space_after = Pt(24)
    run = p.add_run("_" * 60)
    set_font(run, size=10, color=RGBColor(0xBF, 0xBF, 0xBF))

    # Client / project
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Gent Machine — Davenport Maintenance Assistant")
    set_font(run, size=14, color=DARK_GRAY)

    # Blank space
    for _ in range(3):
        doc.add_paragraph().space_after = Pt(0)

    # Prepared by
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Prepared by: Aidan Systems")
    set_font(run, size=12, color=DARK_GRAY)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("(subcontractor to MAGNET)")
    set_font(run, size=11, color=RGBColor(0x66, 0x66, 0x66))

    doc.add_paragraph().space_after = Pt(0)

    # Date
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Date: March 2026")
    set_font(run, size=12, color=DARK_GRAY)

    doc.add_paragraph().space_after = Pt(0)

    # Status
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Status: DRAFT")
    set_font(run, size=12, bold=True, color=RGBColor(0xC0, 0x39, 0x2B))

    # Page break after title
    doc.add_page_break()


def add_heading(doc, text, level=1):
    """Add a heading with consistent styling."""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.name = "Calibri"
        run.font.color.rgb = NAVY
    return heading


def add_body(doc, text):
    """Add body text with bold/italic parsing."""
    p = doc.add_paragraph()
    add_formatted_text(p, text)
    return p


def add_italic_body(doc, text):
    """Add italic body text."""
    p = doc.add_paragraph()
    # Strip leading/trailing * if present
    clean = text.strip('*').strip()
    run = p.add_run(clean)
    set_font(run, italic=True, color=RGBColor(0x55, 0x55, 0x55))
    return p


def add_bullet(doc, text):
    """Add a bullet point with formatted text."""
    p = doc.add_paragraph(style='List Bullet')
    # Clear default run, use formatted text
    p.clear()
    add_formatted_text(p, text)
    return p


def build_document():
    """Build the complete SOW document."""
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    font.color.rgb = DARK_GRAY

    # Set margins to 1 inch
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Set heading styles
    for level in range(1, 4):
        style_name = f'Heading {level}'
        if style_name in doc.styles:
            h_style = doc.styles[style_name]
            h_style.font.name = 'Calibri'
            h_style.font.color.rgb = NAVY
            if level == 1:
                h_style.font.size = Pt(18)
                h_style.paragraph_format.space_before = Pt(24)
                h_style.paragraph_format.space_after = Pt(12)
            elif level == 2:
                h_style.font.size = Pt(15)
                h_style.paragraph_format.space_before = Pt(18)
                h_style.paragraph_format.space_after = Pt(8)
            elif level == 3:
                h_style.font.size = Pt(13)
                h_style.paragraph_format.space_before = Pt(14)
                h_style.paragraph_format.space_after = Pt(6)

    # ===== TITLE PAGE =====
    create_title_page(doc)

    # ===== BACKGROUND =====
    add_heading(doc, "Background", level=1)
    add_body(doc, "Phase 1 delivered a production AI maintenance assistant for Davenport Model B screw machines at Gent Machine's Cleveland facility. The system is live — 1,241 indexed documents, a 570-vertex machine knowledge graph, and a shop floor web application with built-in feedback and analytics.")
    add_body(doc, "On March 18, 2026, a Gent machinist used the system for real troubleshooting — asking about flat thread peaks, thrust bearing replacement, cutoff steps, and short parts. The system delivered cited, step-by-step guidance. It also revealed knowledge gaps: questions about extension pins and specific brand comparisons weren't in the knowledge base.")
    add_body(doc, "**The system works. It needs more knowledge and more pilots.**")
    add_body(doc, "Phase 2 transitions from **Configure & Deploy** to **Train & Tune** — getting every machinist into the cockpit, expanding the knowledge base, and coaching the team through adoption into daily reliance.")

    doc.add_page_break()

    # ===== THE FLIGHT PATH =====
    add_heading(doc, "The Flight Path — Where Gent Is Now", level=1)
    add_body(doc, "Aidan's engagement model follows six stages. Each has a flight term (the brand) and a business term (the conversation):")

    add_table(doc,
        ["Stage", "Flight", "Business", "What's Happening"],
        [
            ["0", "**PRE-FLIGHT**", "**Assess**", "Evaluate the operation, build the roadmap"],
            ["1", "**CONFIGURE**", "**Configure & Deploy**", "Configure the AI platform for this operation and deploy to the shop floor"],
            ["2", "**TAXI**", "**Train & Tune**", "Train the team, tune the agent to their language, validate with real questions"],
            ["3", "**WHEELS UP**", "**Adopt**", "Team is using it and telling us what's missing"],
            ["4", "**CLIMBING**", "**Expand**", "More knowledge, more users, more value"],
            ["5", "**CRUISING**", "**Compound**", "Daily operations run on it — every month it gets smarter"],
        ]
    )

    add_heading(doc, "Gent's Current Position", level=2)
    add_table(doc,
        ["Stage", "Status"],
        [
            ["PRE-FLIGHT / Assess", "Complete (Phase 1 scoping)"],
            ["CONFIGURE / Configure & Deploy", "Complete (Phase 1 delivery)"],
            ["TAXI / Train & Tune", "**Complete** — Aidan conducted onsite training with the team. Machinists have used the system for real troubleshooting (3/18 session: thread peaks, thrust bearings, cutoff steps, short parts)."],
            ["WHEELS UP / Adopt", "**In progress** — System is being used but not yet a daily habit. Knowledge gaps being identified. Dave (retired expert) knowledge captured but team needs more content to build trust."],
            ["CLIMBING / Expand", "Future"],
            ["CRUISING / Compound", "Future"],
        ]
    )

    add_body(doc, "**Phase 2 picks up at WHEELS UP** — driving adoption from occasional use to daily habit, expanding the knowledge base, and building toward the growth goal: enabling Gent to staff and run more Davenports.")

    doc.add_page_break()

    # ===== INDUSTRIAL PILOT FRAMEWORK =====
    add_heading(doc, "The Industrial Pilot Framework", level=1)
    add_body(doc, "Drawing on the **Industrial Athlete Operating System** (Lamoncha & Figley, NAM Manufacturer of the Year — Ohio manufacturing), extended by Aidan into the cockpit:")

    add_table(doc,
        ["Traditional", "Industrial Athlete", "Industrial Pilot (Aidan)"],
        [
            ["**Manager** — controls", "**Coach** — develops", "**Flight Instructor** — teaches them to fly solo"],
            ["**Employee** — does tasks", "**Athlete** — performs", "**Pilot** — navigates independently with AI"],
            ["**Workplace** — clock in", "**Performance Center** — excel", "**Cockpit** — AI instruments, checklists, navigation"],
            ["**Egosystem** — silos", "**Ecosystem** — collaboration", "**Airspace** — shared knowledge, governed, connected"],
            ["**Time Card**", "**Visual Earnings** — scoreboard", "**Flight Instruments + Altitude Report** — performance visibility"],
        ]
    )

    add_heading(doc, "Roles", level=2)
    add_table(doc,
        ["Role", "Who", "What They Do"],
        [
            ["**Flight Instructor**", "Aidan (Managed AIR)", "Coaches adoption, diagnoses gaps, calibrates instruments, teaches the team to fly"],
            ["**Pilot**", "Machinist", "Navigates with AI instruments, reports what's working and what's not"],
            ["**Co-Pilot**", "The AI system", "Always present, provides guidance and citations — but the human makes the calls"],
            ["**Ground Crew**", "Senior machinists / Rich", "Contribute operational knowledge, validate content, identify what's missing"],
        ]
    )

    p = doc.add_paragraph()
    run = p.add_run("Note: Dave (Gent's longtime expert) has retired. His knowledge was captured in Phase 1. The system now carries that expertise — the Ground Crew role shifts to Gent's current senior operators and leadership.")
    set_font(run, italic=True, color=RGBColor(0x55, 0x55, 0x55))

    doc.add_page_break()

    # ===== PHASE 2 PLAN =====
    add_heading(doc, "Phase 2 Plan — Stage by Stage", level=1)

    add_heading(doc, "Completed: TAXI / Train & Tune", level=2)
    add_body(doc, "Aidan conducted onsite training with Gent's team. Machinists used the system for real troubleshooting. Knowledge gaps were identified (extension pin, brand comparisons, oil specifications). The team knows how to use the cockpit.")

    doc.add_page_break()

    # ===== MONTHS 1-3: WHEELS UP =====
    add_heading(doc, "Months 1-3: WHEELS UP / Adopt", level=2)
    add_italic_body(doc, '"It\'s where they go first — not last."')
    add_body(doc, 'The system works. The team has been trained. The challenge now is **habit formation** — moving from "I\'ll try it when I remember" to "I check the cockpit before I do anything else." This requires three things: the system needs to answer well enough that it earns trust, the team needs visibility into how it\'s helping, and leadership needs to set the expectation.')

    add_heading(doc, "Making the Cockpit Unavoidable", level=3)
    add_table(doc,
        ["Activity", "Description", "Who"],
        [
            ['"Ask the Cockpit First" — as a team expectation', "Rich tells his team: before you escalate, check the system. If it helps, great. If it doesn't, flag it — that's how we make it better. This isn't a suggestion; it's how we work now.", "Rich (with Aidan coaching)"],
            ["**Shop floor display board**", "Printed weekly — posted where everyone sees it. Shows: queries this week, best Q&A, knowledge gap closed, team stats. Makes the system visible even when nobody's at the terminal.", "Aidan prepares, Gent posts"],
            ["**Fix the known gaps — fast**", "Extension pin, oil specifications, thrust bearing brands, and any other gaps from the 3/18 session. If a machinist tries the system and hits a gap they already reported, trust erodes. Close these first.", "Aidan (Flight Instructor)"],
        ]
    )

    add_heading(doc, "Building Trust Through Quality", level=3)
    add_table(doc,
        ["Activity", "Description", "Who"],
        [
            ["**Content sprint — top 10 scenarios**", "Work with Rich and senior machinists to identify the 10 most common troubleshooting scenarios. Verify the system handles each one well. Fill gaps.", "Aidan + Gent senior operators"],
            ["**Agent tuning for Gent's language**", 'Real machinists misspell things ("thrush bearim"), use shorthand, and describe problems differently than manuals do. Tune the agent to handle how Gent\'s team actually talks.', "Aidan (Flight Instructor)"],
            ['"Did this help?" follow-up', "When a machinist uses the system, check in casually: \"Did that answer work?\" This isn't surveillance — it's coaching. Their feedback identifies gaps faster than the thumbs-up button.", "Rich / shift leads"],
        ]
    )

    add_heading(doc, "Making Contribution Easy", level=3)
    add_table(doc,
        ["Activity", "Description", "Who"],
        [
            ["**Flag + notes — type what the system missed**", 'When a machinist gets a bad or incomplete answer, flag it and type a short note: "extension pin was the real fix" or "wrong torque spec." Even a few words helps Aidan close the gap.', "All machinists"],
            ["**Dave reviews on his phone**", "Dave is retired but still accessible. He can use the app on his phone to review answers, flag what's wrong, and add corrections. Aidan also conducts periodic interviews with Dave to capture specific knowledge areas.", "Dave (remote, on his schedule)"],
            ['"What would Dave have said?"', "Show senior machinists the system's answer to a real question. Ask: \"Is this right? What's missing?\" Their corrections go straight into the knowledge base.", "Aidan + senior machinists"],
            ["**Flag = contribution, not complaint**", "Reframe flagging a bad answer as helping the team, not criticizing the system. Every flag makes the cockpit smarter for everyone. Recognize it.", "Rich reinforces the culture"],
        ]
    )

    add_heading(doc, "Coaching Cadence", level=3)
    add_table(doc,
        ["Activity", "Frequency", "Description"],
        [
            ["**Biweekly check-in with Rich**", "Every 2 weeks", '15 minutes. "Here\'s what your team asked. Here\'s what worked. Here\'s what we fixed. Here\'s what\'s still missing." Not a status call — a coaching session.'],
            ["**Monthly Altitude Report**", "Monthly", "Posted on the shop floor AND sent to Rich. Adoption progress, satisfaction, gaps closed, team contributions. Framed as team achievement."],
            ["**Knowledge gap triage**", "Weekly", "Aidan reviews flags, thumbs-down, and typed notes. Prioritizes and closes gaps. The faster gaps close, the faster trust builds."],
        ]
    )

    # Targets
    add_heading(doc, "Targets", level=3)
    add_bullet(doc, "2+ queries per machinist per week by end of month 2")
    add_bullet(doc, "70%+ of the top 10 troubleshooting scenarios well-covered in the KB")
    add_bullet(doc, "Feedback flowing regularly (flags, thumbs, typed notes)")
    add_bullet(doc, 'Rich actively reinforcing "Ask the Cockpit First"')

    add_body(doc, '**What success looks like at the end of WHEELS UP:** A machinist has a problem → they walk to the terminal first → they get a useful answer most of the time → when they don\'t, they flag it and it gets fixed within days. The cockpit is becoming the default, not the afterthought.')

    doc.add_page_break()

    # ===== MONTHS 4-6: CLIMBING =====
    add_heading(doc, "Months 4-6: CLIMBING / Expand", level=2)
    add_italic_body(doc, '"The system is earning trust — and enabling growth."')
    add_body(doc, "By now the knowledge base covers common scenarios, usage is regular, and the team trusts the cockpit. This is where the growth story begins.")

    add_table(doc,
        ["Activity", "Description", "Who"],
        [
            ["**Add setup procedures to the KB**", "Not just troubleshooting — setup, changeover, and adjustment procedures. This is what enables a less experienced operator to set up a machine.", "Aidan + Gent senior operators"],
            ["**Flight Instruments live**", "Leaderboard, personal stats, team progress — visible in the system. Points for queries (1), feedback (2), flags that become content (5).", "Aidan builds, team uses"],
            ["**Content sprint — deep areas**", "Systematic expansion into areas the top 10 didn't cover. Driven by what machinists are actually asking and flagging.", "Aidan (Flight Instructor)"],
            ["**Graph ontology expansion**", "New knowledge graph nodes and edges for areas machinists are asking about — improves search relevance.", "Aidan (Flight Instructor)"],
            ["**Monthly Altitude Report**", 'Team achievement focus — adoption curve, knowledge growth, "the system answered X more questions this month than last."', "Aidan prepares, posted on floor"],
            ["**Growth planning**", "Begin the conversation: which additional Davenport jobs could a junior hire handle with cockpit support? What setup content would they need?", "Rich + Aidan"],
            ["**Expansion scoping**", "If the growth path is clear, scope the work center expansion ($10k) to add new job types or machine configurations.", "Aidan + MAGNET"],
        ]
    )

    add_heading(doc, "Targets", level=3)
    add_bullet(doc, "5+ queries per week per machinist")
    add_bullet(doc, "70%+ satisfaction rate")
    add_bullet(doc, "Setup procedures indexed for top job types")
    add_bullet(doc, "Growth conversation underway — path to additional Davenport(s) identified")

    doc.add_page_break()

    # ===== MONTH 7+: CRUISING =====
    add_heading(doc, "Month 7+: CRUISING / Compound", level=2)
    add_italic_body(doc, '"It runs. It compounds. It grows the business."')
    add_body(doc, "The system is comprehensive for Davenport operations. Usage is daily. The team trusts it. Now it's about maintaining altitude and enabling growth.")

    add_table(doc,
        ["Activity", "Description"],
        [
            ["**Feedback triage**", "Respond to flags and thumbs-down, close gaps as they surface"],
            ["**Monthly Altitude Report**", "Usage trends, satisfaction, cost, recommendations"],
            ["**Infrastructure health**", "Azure monitoring, security, SOC 2 compliance"],
            ["**Minor adjustments**", "Jargon additions, instruction tweaks, content updates"],
            ["**New hire onboarding**", "When Gent hires a junior machinist, the cockpit is part of their day-one toolkit. Flight Instructor monitors their query patterns to identify gaps."],
            ["**Expansion execution**", "Configure new work centers ($10k each) as Gent grows into additional capacity"],
        ]
    )

    add_heading(doc, "Targets", level=3)
    add_bullet(doc, "80%+ weekly active machinists")
    add_bullet(doc, "75%+ satisfaction rate")
    add_bullet(doc, "New hires productive on basic jobs within weeks, not months")
    add_bullet(doc, "Additional Davenport(s) running with cockpit-supported operators")

    doc.add_page_break()

    # ===== SERVICE TIERS =====
    add_heading(doc, "Service Tiers", level=1)
    add_body(doc, "All tiers include Azure infrastructure management and SOC 2 compliance. Rate is **$200/hour** across all tiers — tiers determine included hours, not rate.")

    # Small
    add_heading(doc, "Small — Maintain Altitude", level=2)
    add_body(doc, "**$800/month**")
    add_table(doc,
        ["Included", "Hours"],
        [
            ["Azure infrastructure: hosting, monitoring, security, patching", "—"],
            ["SOC 2 compliance controls and evidence", "—"],
            ["Feedback triage — review flags from the cockpit", "0.5 hr"],
            ["Monthly Altitude Report", "0.5 hr"],
            ["Minor flight instrument adjustments", "1 hr"],
            ["**Total**", "**2 hrs**"],
        ]
    )
    add_body(doc, "Best for: CRUISING stage — system healthy, minimal gaps.")

    # Medium
    add_heading(doc, "Medium — Climb", level=2)
    add_body(doc, "**$1,200/month**")
    add_table(doc,
        ["Included", "Hours"],
        [
            ["Everything in Small", "2 hrs"],
            ["Proactive knowledge gap analysis", "1 hr"],
            ["Content creation — 1-2 new KB articles with SME", "1.5 hrs"],
            ["Agent tuning and flight instrument calibration", "0.5 hr"],
            ["Monthly coaching call with Gent stakeholders", "1 hr"],
            ["**Total**", "**6 hrs**"],
        ]
    )
    add_body(doc, "Best for: TAXI through CLIMBING — building the knowledge base and coaching adoption.")

    # Large
    add_heading(doc, "Large — Full Throttle", level=2)
    add_body(doc, "**$2,000/month**")
    add_table(doc,
        ["Included", "Hours"],
        [
            ["Everything in Medium", "6 hrs"],
            ["Content sprint — 3-5 new KB articles or knowledge area buildout", "2 hrs"],
            ["Search quality audit and retrieval tuning", "1 hr"],
            ["Graph ontology expansion", "1 hr"],
            ["**Total**", "**10 hrs**"],
        ]
    )
    add_body(doc, "Best for: Major improvement months or preparing for a new work center.")

    # Overage
    add_heading(doc, "Overage", level=2)
    add_body(doc, "Additional hours beyond any tier: **$200/hour**, billed monthly in arrears.")

    doc.add_page_break()

    # ===== WORK CENTER EXPANSION =====
    add_heading(doc, "Work Center Expansion", level=1)
    add_table(doc,
        ["Item", "Fee"],
        [
            ["**New work center** (new machine type or operational area)", "**$10,000**"],
        ]
    )
    add_body(doc, "This is a **CONFIGURE / Configure & Deploy** for a new route on the existing platform. Includes knowledge capture, index expansion, agent reconfiguration, ontology extension, and verification.")
    add_body(doc, "**Platform incentive:** Monthly Managed AIR fee covers ALL work centers. More routes in the airspace, more value per dollar.")

    add_table(doc,
        ["Work Centers", "Monthly (Medium)", "Annual Cost per Work Center"],
        [
            ["1", "$1,200", "$14,400"],
            ["2", "$1,200", "$7,200"],
            ["3", "$1,200", "$4,800"],
        ]
    )

    doc.add_page_break()

    # ===== FLIGHT INSTRUMENTS =====
    add_heading(doc, "Flight Instruments & Altitude Report", level=1)

    add_heading(doc, "Always-On Flight Instruments", level=2)
    add_bullet(doc, "Query volume and user activity (JSONL analytics lake)")
    add_bullet(doc, "Satisfaction signals (feedback table — thumbs up/down/flag)")
    add_bullet(doc, "Response performance (timing breakdowns)")
    add_bullet(doc, "Knowledge graph utilization (edge hit counters)")
    add_bullet(doc, "Admin dashboard with real-time stats and conversation viewer")

    add_heading(doc, "Monthly Altitude Report (all tiers)", level=2)
    bullets = [
        "Adoption stage progress — where are we on the flight path?",
        "Query volume trend vs. prior month",
        "Active pilots and new pilots this month",
        "Satisfaction breakdown and knowledge gap inventory",
        "Content improvements delivered",
        "Azure cost actuals",
        "Recommendations for next month",
    ]
    for i, b in enumerate(bullets):
        p = doc.add_paragraph(style='List Bullet')
        p.clear()
        run = p.add_run(f"{i+1}. {b}")
        set_font(run)

    doc.add_page_break()

    # ===== AZURE INFRASTRUCTURE =====
    add_heading(doc, "Azure Infrastructure", level=1)
    add_body(doc, "Current monthly cost: **~$130/month**")

    add_table(doc,
        ["Service", "Cost", "Purpose"],
        [
            ["Azure Cognitive Search (Basic)", "~$70", "Unified index, 1,241 docs"],
            ["Azure Container Apps", "~$24", "Foundry agent runtime"],
            ["Azure App Service", "~$16", "Web app hosting"],
            ["Container Registry", "~$5", "Container images"],
            ["Foundry Models (AOAI)", "~$3-27", "GPT-5-mini + embeddings (usage-based)"],
            ["Storage + Tables", "~$0.25", "Blobs, feedback, analytics, verification ledger"],
            ["Cosmos DB (Serverless)", "~$0.04", "Graph ontology"],
            ["Functions (Flex Consumption)", "$0.00", "API layer"],
        ]
    )
    add_body(doc, "Included in all tiers. Covers management, monitoring, security, SOC 2 compliance.")

    doc.add_page_break()

    # ===== RECOMMENDED PLAN =====
    add_heading(doc, "Recommended Plan & Investment", level=1)
    add_table(doc,
        ["Period", "Tier", "Stage", "Monthly (Aidan)", "Focus"],
        [
            ["Months 1-3", "Medium", "WHEELS UP / Adopt", "$1,200", "Drive daily use, close gaps, build trust, coaching cadence"],
            ["Months 4-6", "Medium", "CLIMBING / Expand", "$1,200", "Add setup procedures, Flight Instruments, growth planning"],
            ["Month 7+", "Small", "CRUISING / Compound", "$800", "Maintain, support new hires, enable expansion"],
        ]
    )

    add_body(doc, "**Year 1 estimate (Aidan fees):** $12,000 (6 x $1,200 + 6 x $800)")
    p = doc.add_paragraph()
    run = p.add_run("Note: All fees are Aidan's rates to MAGNET. MAGNET applies their own markup to Gent.")
    set_font(run, italic=True, color=RGBColor(0x55, 0x55, 0x55))

    doc.add_page_break()

    # ===== TERM AND INVOICING =====
    add_heading(doc, "Term and Invoicing", level=1)
    add_bullet(doc, "**Initial term:** 3 months (Medium tier)")
    add_bullet(doc, "**After initial term:** Month-to-month, 15-day notice to change tier or cancel")
    add_bullet(doc, "**Invoicing:** Monthly in advance (tier) + monthly in arrears (overage at $200/hr)")
    add_bullet(doc, "**Work center expansions:** 50% at kickoff / 50% at completion")
    add_bullet(doc, "**Azure costs:** Included in all tiers")

    doc.add_page_break()

    # ===== WHAT THIS DOES NOT INCLUDE =====
    add_heading(doc, "What This Does NOT Include", level=1)
    add_bullet(doc, "Net-new application development beyond cockpit enhancements")
    add_bullet(doc, "Migration to a different Azure tenant")
    add_bullet(doc, "Teams or Copilot integration")
    add_bullet(doc, "Image/drawing recognition")
    add_bullet(doc, "Advanced enterprise security (Private Link, complex RBAC)")
    add_body(doc, "These can be scoped as expansion projects or a future phase.")

    doc.add_page_break()

    # ===== SUMMARY =====
    add_heading(doc, "Summary", level=1)
    add_table(doc,
        ["Item", "Aidan Fee"],
        [
            ["**Small** — Maintain Altitude (2 hrs)", "$800/month"],
            ["**Medium** — Climb (6 hrs)", "$1,200/month"],
            ["**Large** — Full Throttle (10 hrs)", "$2,000/month"],
            ["Additional hours (any tier)", "$200/hour"],
            ["New work center expansion", "$10,000"],
            ["Initial commitment", "3 months"],
        ]
    )

    # Add page numbers
    add_page_number(doc)

    # Save
    doc.save(OUTPUT_PATH)
    print(f"Document saved to: {OUTPUT_PATH}")


if __name__ == "__main__":
    build_document()
