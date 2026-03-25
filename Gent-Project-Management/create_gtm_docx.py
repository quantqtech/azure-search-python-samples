"""
Convert the MAGNET GTM markdown to a professional Word document.
Uses python-docx to create polished formatting for a sales/partnership document.
"""

import sys
import os
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import re


# -- Color constants --
NAVY = RGBColor(0x1B, 0x2A, 0x4A)       # Dark navy for headers
DARK_BLUE = RGBColor(0x1F, 0x3A, 0x5F)   # Table header background
LIGHT_GRAY = RGBColor(0xF2, 0xF2, 0xF2)  # Alternating row shading
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
BLACK = RGBColor(0x00, 0x00, 0x00)
MEDIUM_GRAY = RGBColor(0x66, 0x66, 0x66)
ACCENT_BLUE = RGBColor(0x2E, 0x75, 0xB6)  # For subtle accents


def set_cell_shading(cell, color_hex):
    """Apply background shading to a table cell."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_borders(cell, top=None, bottom=None, left=None, right=None):
    """Set borders on a table cell. Each border arg is (size_eighths, color_hex, style)."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}/>')

    for side, val in [("top", top), ("bottom", bottom), ("left", left), ("right", right)]:
        if val:
            sz, color, style = val
            border_el = parse_xml(
                f'<w:{side} {nsdecls("w")} w:val="{style}" w:sz="{sz}" '
                f'w:space="0" w:color="{color}"/>'
            )
            tcBorders.append(border_el)

    tcPr.append(tcBorders)


def set_table_borders(table):
    """Apply thin borders to the entire table."""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="4" w:space="0" w:color="BFBFBF"/>'
        f'  <w:left w:val="single" w:sz="4" w:space="0" w:color="BFBFBF"/>'
        f'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="BFBFBF"/>'
        f'  <w:right w:val="single" w:sz="4" w:space="0" w:color="BFBFBF"/>'
        f'  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="BFBFBF"/>'
        f'  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="BFBFBF"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)


def add_formatted_text(paragraph, text, font_name="Calibri", font_size=11, bold=False, color=None, italic=False):
    """Add a run of text with formatting to a paragraph."""
    run = paragraph.add_run(text)
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color
    return run


def parse_inline_formatting(paragraph, text, font_name="Calibri", font_size=11, color=None):
    """Parse markdown bold (**text**) and render as formatted runs."""
    # Split on bold markers
    parts = re.split(r'\*\*(.+?)\*\*', text)
    for i, part in enumerate(parts):
        if not part:
            continue
        is_bold = (i % 2 == 1)  # Odd indices are inside **...**
        add_formatted_text(paragraph, part, font_name, font_size, bold=is_bold, color=color)


def add_page_number(doc):
    """Add page numbers to the footer."""
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Page number field
        run = p.add_run()
        run.font.name = "Calibri"
        run.font.size = Pt(9)
        run.font.color.rgb = MEDIUM_GRAY

        fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
        run._r.append(fldChar1)

        instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
        run._r.append(instrText)

        fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
        run._r.append(fldChar2)


def create_title_page(doc):
    """Create a professional title page."""
    # Add some blank space at top
    for _ in range(6):
        p = doc.add_paragraph()
        p.space_after = Pt(0)
        p.space_before = Pt(0)

    # Main title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.space_after = Pt(6)
    add_formatted_text(p, "The Industrial Pilot Platform", "Calibri", 32, bold=True, color=NAVY)

    # Subtitle
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.space_after = Pt(24)
    add_formatted_text(p, "AI-Powered Operational Intelligence for Ohio Manufacturers", "Calibri", 16, color=ACCENT_BLUE, italic=True)

    # Horizontal line via a thin table
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("_" * 60)
    run.font.color.rgb = RGBColor(0xBF, 0xBF, 0xBF)
    run.font.size = Pt(10)

    # Spacer
    doc.add_paragraph().space_after = Pt(12)

    # MAGNET Go-To-Market
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.space_after = Pt(36)
    add_formatted_text(p, "MAGNET Go-To-Market", "Calibri", 18, bold=True, color=NAVY)

    # Prepared by
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.space_after = Pt(6)
    add_formatted_text(p, "Prepared by: ", "Calibri", 12, color=MEDIUM_GRAY)
    add_formatted_text(p, "Aidan Systems", "Calibri", 12, bold=True, color=NAVY)

    # Date
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.space_after = Pt(6)
    add_formatted_text(p, "Date: March 2026", "Calibri", 12, color=MEDIUM_GRAY)

    # Status
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.space_after = Pt(6)
    add_formatted_text(p, "Status: ", "Calibri", 12, color=MEDIUM_GRAY)
    add_formatted_text(p, "DRAFT", "Calibri", 12, bold=True, color=RGBColor(0xC0, 0x39, 0x2B))
    add_formatted_text(p, " \u2014 for MAGNET partnership discussion", "Calibri", 12, color=MEDIUM_GRAY)

    # Reference customer
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.space_after = Pt(6)
    add_formatted_text(p, "Reference Customer: ", "Calibri", 12, color=MEDIUM_GRAY)
    add_formatted_text(p, "Gent Machine", "Calibri", 12, bold=True, color=NAVY)
    add_formatted_text(p, " (Davenport screw machines \u2014 live, WHEELS UP)", "Calibri", 12, color=MEDIUM_GRAY)

    # Page break after title page
    doc.add_page_break()


def add_table_from_rows(doc, headers, rows):
    """Create a formatted table with header shading and alternating rows."""
    num_cols = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=num_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    # Set table borders
    set_table_borders(table)

    # Header row
    header_row = table.rows[0]
    for i, header_text in enumerate(headers):
        cell = header_row.cells[i]
        set_cell_shading(cell, "1F3A5F")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        # Parse bold in headers too
        clean_text = re.sub(r'\*\*(.+?)\*\*', r'\1', header_text)
        add_formatted_text(p, clean_text, "Calibri", 10, bold=True, color=WHITE)

    # Data rows
    for row_idx, row_data in enumerate(rows):
        row = table.rows[row_idx + 1]
        for col_idx in range(num_cols):
            cell = row.cells[col_idx]
            # Alternating row shading
            if row_idx % 2 == 1:
                set_cell_shading(cell, "F2F2F2")

            cell_text = row_data[col_idx] if col_idx < len(row_data) else ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            parse_inline_formatting(p, cell_text, "Calibri", 10, color=BLACK)

    # Add spacing after table
    doc.add_paragraph().space_after = Pt(6)

    return table


def parse_table_block(lines):
    """Parse a markdown table block into headers and rows."""
    # Filter out separator lines (|---|---|)
    content_lines = [l for l in lines if not re.match(r'^\|[\s\-|:]+\|$', l)]

    if not content_lines:
        return [], []

    def split_row(line):
        # Split by | and strip whitespace, ignoring empty first/last from leading/trailing |
        cells = [c.strip() for c in line.split('|')]
        # Remove empty strings from leading/trailing pipes
        if cells and cells[0] == '':
            cells = cells[1:]
        if cells and cells[-1] == '':
            cells = cells[:-1]
        return cells

    headers = split_row(content_lines[0])
    rows = [split_row(l) for l in content_lines[1:]]
    return headers, rows


def process_markdown(doc, md_text):
    """Parse the markdown content and add it to the document."""
    lines = md_text.split('\n')
    i = 0
    # Skip the title block (first 8 lines are title, subtitle, metadata)
    # We handle those in the title page
    # Find the first --- separator after the metadata
    skip_until_first_section = True

    while i < len(lines):
        line = lines[i]

        # Skip the metadata block at the top (title, subtitle, prepared by, etc.)
        if skip_until_first_section:
            if line.strip() == '---':
                # Found the first separator after metadata
                skip_until_first_section = False
                i += 1
                # Skip the next --- too if it's the first real content separator
                continue
            i += 1
            continue

        # Skip horizontal rules
        if line.strip() == '---':
            i += 1
            continue

        # Heading 1 (##)
        if line.startswith('## '):
            # Page break before major sections
            doc.add_page_break()
            heading_text = line[3:].strip()
            p = doc.add_heading(heading_text, level=1)
            # Style the heading
            for run in p.runs:
                run.font.name = "Calibri"
                run.font.color.rgb = NAVY
            i += 1
            continue

        # Heading 2 (###)
        if line.startswith('### ') and not line.startswith('#### '):
            heading_text = line[4:].strip()
            p = doc.add_heading(heading_text, level=2)
            for run in p.runs:
                run.font.name = "Calibri"
                run.font.color.rgb = NAVY
            i += 1
            continue

        # Heading 3 (####)
        if line.startswith('#### '):
            heading_text = line[5:].strip()
            p = doc.add_heading(heading_text, level=3)
            for run in p.runs:
                run.font.name = "Calibri"
                run.font.color.rgb = NAVY
            i += 1
            continue

        # Table block
        if line.strip().startswith('|') and i + 1 < len(lines):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_lines.append(lines[i].strip())
                i += 1
            headers, rows = parse_table_block(table_lines)
            if headers:
                add_table_from_rows(doc, headers, rows)
            continue

        # Blockquote
        if line.strip().startswith('>'):
            quote_lines = []
            while i < len(lines) and (lines[i].strip().startswith('>') or lines[i].strip() == ''):
                stripped = lines[i].strip()
                if stripped == '' and (i + 1 >= len(lines) or not lines[i + 1].strip().startswith('>')):
                    break
                if stripped.startswith('>'):
                    stripped = stripped[1:].strip()
                quote_lines.append(stripped)
                i += 1

            quote_text = '\n'.join(quote_lines)

            # Create a styled paragraph for the blockquote
            for q_line in quote_lines:
                if q_line.strip() == '':
                    doc.add_paragraph().space_after = Pt(2)
                    continue
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.5)
                p.paragraph_format.space_after = Pt(4)
                # Add a left border via indentation + italic styling
                parse_inline_formatting(p, q_line, "Calibri", 11, color=RGBColor(0x44, 0x44, 0x44))
                for run in p.runs:
                    run.font.italic = True
            continue

        # Bullet points
        if line.strip().startswith('- ') or line.strip().startswith('* '):
            bullet_text = line.strip()[2:]
            p = doc.add_paragraph(style='List Bullet')
            parse_inline_formatting(p, bullet_text, "Calibri", 11, color=BLACK)
            i += 1
            continue

        # Numbered list
        match = re.match(r'^(\d+)\.\s+(.+)', line.strip())
        if match:
            item_text = match.group(2)
            p = doc.add_paragraph(style='List Number')
            parse_inline_formatting(p, item_text, "Calibri", 11, color=BLACK)
            i += 1
            continue

        # Regular paragraph
        if line.strip():
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(6)
            parse_inline_formatting(p, line.strip(), "Calibri", 11, color=BLACK)
            i += 1
            continue

        # Empty line - skip
        i += 1


def setup_styles(doc):
    """Configure document styles for consistent formatting."""
    style = doc.styles['Normal']
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)
    font.color.rgb = BLACK

    # Heading 1
    h1 = doc.styles['Heading 1']
    h1.font.name = "Calibri"
    h1.font.size = Pt(22)
    h1.font.color.rgb = NAVY
    h1.font.bold = True
    h1.paragraph_format.space_before = Pt(12)
    h1.paragraph_format.space_after = Pt(8)

    # Heading 2
    h2 = doc.styles['Heading 2']
    h2.font.name = "Calibri"
    h2.font.size = Pt(16)
    h2.font.color.rgb = NAVY
    h2.font.bold = True
    h2.paragraph_format.space_before = Pt(10)
    h2.paragraph_format.space_after = Pt(6)

    # Heading 3
    h3 = doc.styles['Heading 3']
    h3.font.name = "Calibri"
    h3.font.size = Pt(13)
    h3.font.color.rgb = NAVY
    h3.font.bold = True
    h3.paragraph_format.space_before = Pt(8)
    h3.paragraph_format.space_after = Pt(4)


def main():
    # Read the markdown source
    md_path = r"c:\repos\my-projects\azure-search-python-samples\Gent-Project-Management\MAGNET GTM - Industrial Pilot Platform.md"
    out_path = r"c:\repos\my-projects\azure-search-python-samples\Gent-Project-Management\MAGNET GTM - Industrial Pilot Platform.docx"

    with open(md_path, 'r', encoding='utf-8') as f:
        md_text = f.read()

    # Create document
    doc = Document()

    # Set margins to 1 inch on all sides
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Setup styles
    setup_styles(doc)

    # Create title page
    create_title_page(doc)

    # Process the markdown content
    process_markdown(doc, md_text)

    # Add page numbers
    add_page_number(doc)

    # Save
    doc.save(out_path)
    print(f"Document saved to: {out_path}")
    print("Done!")


if __name__ == "__main__":
    main()
